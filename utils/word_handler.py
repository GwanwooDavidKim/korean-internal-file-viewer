# -*- coding: utf-8 -*-
"""
Word 문서 처리 모듈 (Word Document Handler)

python-docx를 사용하여 Word 문서에서 텍스트와 구조 정보를 추출합니다.
"""
from docx import Document
from docx.shared import Inches
import os
from typing import List, Dict, Any, Optional, Tuple


class WordHandler:
    """
    Word 문서 처리를 위한 클래스입니다.
    
    주요 기능:
    - Word 문서 텍스트 추출
    - 문서 구조 분석 (제목, 단락, 표 등)
    - 문서 메타데이터 조회
    - 스타일 정보 추출
    """
    
    def __init__(self):
        """WordHandler 인스턴스를 초기화합니다."""
        self.supported_extensions = ['.docx']  # .doc는 python-docx에서 지원하지 않음
    
    def can_handle(self, file_path: str) -> bool:
        """
        파일이 이 핸들러가 처리할 수 있는 형식인지 확인합니다.
        
        Args:
            file_path (str): 파일 경로
            
        Returns:
            bool: 처리 가능 여부
        """
        return any(file_path.lower().endswith(ext) for ext in self.supported_extensions)
    
    def extract_text(self, file_path: str, include_structure: bool = True) -> str:
        """
        Word 문서에서 텍스트를 추출합니다.
        
        Args:
            file_path (str): Word 파일 경로
            include_structure (bool): 구조 정보 포함 여부
            
        Returns:
            str: 추출된 텍스트
        """
        try:
            doc = Document(file_path)
            text_content = []
            
            if include_structure:
                # 구조를 포함한 텍스트 추출
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        # 스타일 정보 추가
                        style = paragraph.style.name if paragraph.style else "Normal"
                        if "Heading" in style:
                            text_content.append(f"\\n[{style}] {paragraph.text}")
                        else:
                            text_content.append(paragraph.text)
                
                # 표 내용 추출
                for table in doc.tables:
                    text_content.append("\\n=== 표 ===")
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            text_content.append(" | ".join(row_text))
                    text_content.append("=== 표 끝 ===\\n")
            else:
                # 단순 텍스트만 추출
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)
                
                # 표 텍스트 추가
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_content.append(cell.text)
            
            return "\\n".join(text_content)
            
        except Exception as e:
            return f"Word 문서 텍스트 추출 오류: {e}"
    
    def get_document_info(self, file_path: str) -> Dict[str, Any]:
        """
        Word 문서의 상세 정보를 반환합니다.
        
        Args:
            file_path (str): Word 파일 경로
            
        Returns:
            Dict[str, Any]: 문서 정보
        """
        try:
            if not os.path.exists(file_path):
                return {'error': '파일을 찾을 수 없습니다'}
            
            doc = Document(file_path)
            
            # 기본 정보
            file_size = os.path.getsize(file_path)
            
            # 문서 구성 요소 카운트
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            
            # 텍스트가 있는 단락만 카운트
            text_paragraphs = sum(1 for p in doc.paragraphs if p.text.strip())
            
            # 제목 스타일 분석
            heading_styles = {}
            for paragraph in doc.paragraphs:
                if paragraph.style and "Heading" in paragraph.style.name:
                    style_name = paragraph.style.name
                    heading_styles[style_name] = heading_styles.get(style_name, 0) + 1
            
            # 메타데이터 추출
            core_props = doc.core_properties
            
            info = {
                'filename': os.path.basename(file_path),
                'file_size': file_size,
                'file_size_mb': round(file_size / (1024 * 1024), 2),
                'total_paragraphs': paragraph_count,
                'text_paragraphs': text_paragraphs,
                'table_count': table_count,
                'heading_styles': heading_styles,
                'title': core_props.title or '제목 없음',
                'author': core_props.author or '작성자 없음',
                'subject': core_props.subject or '주제 없음',
                'keywords': core_props.keywords or '키워드 없음',
                'created': str(core_props.created) if core_props.created else '생성일 없음',
                'modified': str(core_props.modified) if core_props.modified else '수정일 없음',
                'revision': core_props.revision or '버전 없음',
            }
            
            return info
            
        except Exception as e:
            return {'error': f"문서 정보 조회 오류: {e}"}
    
    def get_document_structure(self, file_path: str) -> List[Dict[str, Any]]:
        """
        문서의 구조 정보를 반환합니다.
        
        Args:
            file_path (str): Word 파일 경로
            
        Returns:
            List[Dict[str, Any]]: 문서 구조 목록
        """
        try:
            doc = Document(file_path)
            structure = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    style_name = paragraph.style.name if paragraph.style else "Normal"
                    
                    element_info = {
                        'index': i,
                        'type': 'paragraph',
                        'style': style_name,
                        'text': paragraph.text[:100] + ('...' if len(paragraph.text) > 100 else ''),
                        'full_text': paragraph.text,
                        'is_heading': "Heading" in style_name,
                    }
                    
                    structure.append(element_info)
            
            # 표 정보 추가
            for i, table in enumerate(doc.tables):
                rows = len(table.rows)
                cols = len(table.columns) if table.rows else 0
                
                # 첫 번째 행의 텍스트 미리보기
                preview_text = ""
                if table.rows:
                    first_row_texts = []
                    for cell in table.rows[0].cells:
                        if cell.text.strip():
                            first_row_texts.append(cell.text.strip())
                    preview_text = " | ".join(first_row_texts[:3])  # 처음 3개 셀만
                
                table_info = {
                    'index': len(structure),
                    'type': 'table',
                    'table_number': i + 1,
                    'rows': rows,
                    'columns': cols,
                    'preview': preview_text[:100] + ('...' if len(preview_text) > 100 else ''),
                }
                
                structure.append(table_info)
            
            return structure
            
        except Exception as e:
            return [{'error': f"문서 구조 분석 오류: {e}"}]
    
    def search_in_document(self, file_path: str, search_term: str, 
                          max_results: int = 20) -> List[Dict[str, Any]]:
        """
        문서에서 특정 텍스트를 검색합니다.
        
        Args:
            file_path (str): Word 파일 경로
            search_term (str): 검색할 텍스트
            max_results (int): 최대 결과 수
            
        Returns:
            List[Dict[str, Any]]: 검색 결과 목록
        """
        try:
            doc = Document(file_path)
            results = []
            search_term = search_term.lower()
            
            # 단락에서 검색
            for i, paragraph in enumerate(doc.paragraphs):
                if search_term in paragraph.text.lower():
                    # 검색어 주변 텍스트 추출
                    text = paragraph.text
                    search_pos = text.lower().find(search_term)
                    
                    # 컨텍스트 생성 (검색어 앞뒤 50자)
                    start = max(0, search_pos - 50)
                    end = min(len(text), search_pos + len(search_term) + 50)
                    context = text[start:end]
                    
                    if start > 0:
                        context = "..." + context
                    if end < len(text):
                        context = context + "..."
                    
                    results.append({
                        'location': f"단락 {i + 1}",
                        'type': 'paragraph',
                        'style': paragraph.style.name if paragraph.style else "Normal",
                        'context': context,
                        'full_text': text,
                    })
                    
                    if len(results) >= max_results:
                        return results
            
            # 표에서 검색
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        if search_term in cell.text.lower():
                            results.append({
                                'location': f"표 {table_idx + 1}, 행 {row_idx + 1}, 열 {col_idx + 1}",
                                'type': 'table_cell',
                                'context': cell.text[:100] + ('...' if len(cell.text) > 100 else ''),
                                'full_text': cell.text,
                            })
                            
                            if len(results) >= max_results:
                                return results
            
            return results
            
        except Exception as e:
            return [{'error': f"문서 검색 오류: {e}"}]